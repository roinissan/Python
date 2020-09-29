import os
import re
from StyleExcel import StyleExcel
from openpyxl import Workbook,load_workbook
import Manager
import docx
import MacroInjection
import win32com.client
import sys
from datetime import datetime
import shutil
#import psutil

class Modifcation:
    def __init__(self,plc_file,graph_file):
        self.plc_file = plc_file
        self.graph_file = graph_file
        self.dir_path = os.path.dirname(plc_file)
        self.plc_file_name = os.path.basename(plc_file).split(".")[0]
        self.graph_file_name = os.path.basename(graph_file).split(".")[0]


    def create_new_folder(self):
        print(self.dir_path)
        print(self.dir_path + "/" + os.path.basename(self.dir_path) + "-" + self.get_process_num())
        try:
            os.makedirs(self.dir_path + "/" + os.path.basename(self.dir_path) + "-" + self.get_process_num())
        except Exception:
            pass


    def get_new_folder_path(self):
        x = self.dir_path + "/" + os.path.basename(self.dir_path) + "-" + self.get_process_num()
        return self.dir_path + "/" + os.path.basename(self.dir_path)+ "-" + self.get_process_num()

    def csv_to_xlsx(self):
        try:
            f = open(self.plc_file, "r", encoding='utf-8')
        except FileNotFoundError:
            print("The file does-not exist")
        csv_file = f.readlines()
        f.close()
        count = 0
        last_lines = []
        num_of_lines = len(csv_file)
        for row in range(num_of_lines-1,-1,-1):
            csv_file[row] = csv_file[row].replace (",",";")
            if (num_of_lines-1 != row and "=" in csv_file[row]):
                count = count +1
                break
            if (row != num_of_lines-1  and count == 0):
                last_lines.append(csv_file[row])
        last_lines.reverse()
        return last_lines


    def get_process_num(self):
        xslx_table = self.csv_to_xlsx()
        process_num  = xslx_table[1].split(";")[3]
        return process_num


    def get_general_info(self):
        xslx_table = self.csv_to_xlsx()
        general_info_list = []
        rev_number = xslx_table[1].split(";")[4]
        data_and_hour = xslx_table[1].split(";")[1]
        program_number = xslx_table[1].split(";")[2]
        hdf_machine_num = xslx_table[1].split(";")[0] + "      rev." + rev_number
        general_info_list.append(data_and_hour)
        general_info_list.append(program_number)
        general_info_list.append(hdf_machine_num)
        return general_info_list


    def get_pn_wo(self):
        xslx_table = self.csv_to_xlsx()
        pn_wo_list =[]
        pn_wo_row = xslx_table[1].strip('\n')
        pn_wo_row = pn_wo_row.split(";")
        pn_wo_list.append( pn_wo_row[5::2])
        pn_wo_list.append(pn_wo_row[6::2])
        return pn_wo_list

    def get_plc_data_table(self):
        xslx_table = self.csv_to_xlsx()
        plc_data = []
        max_temp_c_row = xslx_table[3].strip('\n').split(";")
        hdf_temp_min_max_row =  xslx_table[6].strip('\n').split(";")
        exposure_time_row =  xslx_table[9].strip('\n').split(";")
        vacuum_mecury_row = xslx_table[12].strip('\n').split(";")
        max_temp_c = [max_temp_c_row[2], max_temp_c_row[8], max_temp_c_row[1]]
        hdf_temp_min = [hdf_temp_min_max_row[3], hdf_temp_min_max_row[8], hdf_temp_min_max_row[1]]
        hdf_temp_max = [hdf_temp_min_max_row[4], hdf_temp_min_max_row[8], hdf_temp_min_max_row[2]]
        exposure_time = [exposure_time_row[7], exposure_time_row[8], exposure_time_row[1]]
        vacuum_mecury = [vacuum_mecury_row[2], vacuum_mecury_row[8], vacuum_mecury_row[1]]
        if (len(xslx_table) > 13):
            max_temp_difference_row = xslx_table[13].strip('\n').split(";")
            max_temp_difference = [max_temp_difference_row[2], max_temp_difference_row[8], max_temp_difference_row[1]]
            plc_data.extend((max_temp_c, hdf_temp_min, hdf_temp_max, exposure_time, vacuum_mecury, max_temp_difference))
        else:
            plc_data.extend((max_temp_c,hdf_temp_min,hdf_temp_max,exposure_time,vacuum_mecury))
        return plc_data




    def add_process_num(self,doc):
        doc = doc
        process_num = self.get_process_num()
        for para in doc.paragraphs:
            if("HDF   Process" in  para.text):
                for part in para.runs:
                    style = part.style
                    if ("by controller" in part.text):
                        part.text = process_num
                        part.style = style


    def add_general_info(self,doc):
        doc = doc
        table = doc.tables[0]
        general_info_list = self.get_general_info()
        count_for_info =0
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if (cell.text.find("by controller") != -1 or cell.text.find("xxxxx") != -1):
                        for part in para.runs:
                            style = part.style
                            if ("by controller" in part.text and count_for_info == 0):
                                part.text = general_info_list[0]
                                part.style = style
                                count_for_info = count_for_info +1
                            elif("by controller" in part.text and count_for_info == 1):
                                part.text = general_info_list[1]
                                part.style = style
                            elif(part.text.find("xxxxx") != -1):
                                part.text = general_info_list[2]
                                part.style = style


    def add_pn_wo(self,doc):
        doc = doc
        pn_wo_list = self.get_pn_wo()
        table = doc.tables[1]
        count_for_pn_wo =0
        count_for_pn_wo_list_1,count_for_pn_wo_list_2 = 0,0
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if ((cell.text.find("P.N") != -1) or (cell.text.find("WO")!= -1)):
                        for part in para.runs:
                            style = part.style
                            if ("by controller" in part.text and count_for_pn_wo == 0):
                                if(count_for_pn_wo_list_1 <= len(pn_wo_list[0])-1):
                                    part.text = pn_wo_list[0][count_for_pn_wo_list_1]
                                    part.style = style
                                    count_for_pn_wo = count_for_pn_wo+1
                                    count_for_pn_wo_list_1 = count_for_pn_wo_list_1 +1
                            elif("by controller" in part.text and count_for_pn_wo == 1):
                                if (count_for_pn_wo_list_2 <= len(pn_wo_list[1])-1 ):
                                    part.text = pn_wo_list[1][count_for_pn_wo_list_2]
                                    part.style = style
                                    count_for_pn_wo = 0
                                    count_for_pn_wo_list_2 = count_for_pn_wo_list_2 +1

    def add_plc_data_table(self,doc):
        doc = doc
        plc_data = self.get_plc_data_table()
        table = doc.tables[2]
        row_counter , row_counter_1 = 0,0
        column_counter = 0
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    style = para.style
                    x = cell.text
                    if (para.text.find("by controller") != -1):
                        if (row_counter ==0 and column_counter == 0):
                            para.text = plc_data[0][0]
                            para.style = style
                            column_counter = column_counter +1
                        elif (row_counter ==0 and column_counter == 1):
                            para.text = plc_data[0][1]
                            para.style = style
                            column_counter = 0
                        elif (row_counter ==1 and column_counter == 0):
                            para.text = plc_data[1][0]
                            para.style = style
                            column_counter = column_counter + 1
                        elif (row_counter ==1 and column_counter == 1):
                            para.text = plc_data[1][1]
                            para.style = style
                            column_counter = 0
                        elif (row_counter == 2 and column_counter == 0):
                            para.text = plc_data[2][0]
                            para.style = style
                            column_counter = column_counter + 1
                        elif (row_counter ==2 and column_counter == 1):
                            para.text = plc_data[2][1]
                            para.style = style
                            column_counter = 0
                        elif (row_counter == 3 and column_counter == 0):
                            para.text = plc_data[3][0]
                            para.style = style
                            column_counter = column_counter + 1
                        elif (row_counter ==3 and column_counter == 1):
                            para.text = plc_data[3][1]
                            para.style = style
                            column_counter = 0
                        elif (row_counter == 4 and column_counter == 0):
                            para.text = plc_data[4][0]
                            para.style = style
                            column_counter = column_counter + 1
                        elif (row_counter ==4 and column_counter == 1):
                            para.text = plc_data[4][1]
                            para.style = style
                            column_counter = 0
                        elif (len(plc_data) > 5 and row_counter == 5 and column_counter ==0):
                            para.text = plc_data[5][0]
                            para.style = style
                            column_counter = column_counter + 1
                        elif (len(plc_data) > 5 and row_counter ==5 and column_counter == 1):
                            para.text = plc_data[5][1]
                            para.style = style
                            column_counter = 0
                        if(column_counter == 0):
                            row_counter = row_counter + 1
                    elif(re.search('\d',para.text)):
                        y =re.sub("\d\d",plc_data[0][2],para.text)
                        if (row_counter_1 ==0):
                            para.text = re.sub("\d\d",plc_data[0][2],para.text)
                            para.style = style
                        elif (row_counter_1 ==1):
                            para.text = re.sub("\d\d", plc_data[1][2], para.text)
                            para.style = style
                        elif (row_counter_1 == 2):
                            para.text = re.sub("\d\d", plc_data[2][2], para.text)
                            para.style = style
                        elif (row_counter_1 == 3):
                            para.text = re.sub("\d\d", plc_data[3][2], para.text)
                            para.style = style
                        elif (row_counter_1 == 4):
                            para.text = re.sub("\d\d", plc_data[4][2], para.text)
                            para.style = style
                        elif (len(plc_data) > 5 and row_counter_1 == 5 ):
                            para.text = re.sub("\d\d", plc_data[5][2], para.text)
                            para.style = style
                        row_counter_1 = row_counter_1 + 1

    def modify_datafilewrite(self):
        data_list = tuple(self.csv_to_xlsx())
        new_worksheet = Workbook()
        sheet = new_worksheet.active
        for row in data_list:
            row.replace(",",";")
            sheet.append(row.split(";"))
        new_path = self.get_new_folder_path() + "/" + self.plc_file_name +"-"+ self.get_process_num()+ ".xlsx"
        new_worksheet.save(new_path)
        style = StyleExcel(new_path)
        style.run()

    def write_to_word(self):
        if getattr(sys, 'frozen', False):
            application_path = os.path.dirname(sys.executable)
        elif __file__:
            application_path = os.path.dirname(__file__)
        config_path = os.path.join(application_path, "QA report.docx")
        doc = docx.Document(config_path)
        self.add_pn_wo(doc)
        self.add_process_num(doc)
        self.add_general_info(doc)
        self.add_plc_data_table(doc)
        doc.add_picture(self.get_new_folder_path()+"/chart.png")
        new_path = self.get_new_folder_path() + "/" + self.plc_file_name +"-"+ self.get_process_num() + ".docx"
        doc.save(new_path)


    def diff_times(self,t1, t2):
        try:
            startDate = datetime.strptime(t1.rstrip(' \t\r\n'),"%d/%m/%Y %H:%M:%S")
            try:
                endDate = datetime.strptime(t2.rstrip(' \t\r\n'), "%d/%m/%Y %H:%M:%S")
            except:
                endDate = datetime.strptime(t2.rstrip(' \t\r\n'), "%d/%m/%Y %H:%M")
            difference = (endDate - startDate).total_seconds()/60.0
            #print (startDate, endDate, difference)
            if (difference <= 0 or difference > 100):
                return False
            return True
        except Exception :
            return False

    def diff_times1(self,t1, t2):
        x,y = t1.split()[0], t2.split()[0]
        if(t1.split()[0] != t2.split()[0]):
            return False
        t1,t2 = t1.split()[1].split(":"), t2.split()[1].split(":")
        t1,t2 = t1[:2],t2[:2]
        t1 , t2 = ":".join(t1),":".join(t2)
        ftm = "%H:%M"
        difference = datetime.strptime(t1,ftm) - datetime.strptime(t2,ftm)
        difference = int(str(difference).split(":")[1])
        if(difference >= 0 and difference <100):
            return True
        return False




    def get_graph_data(self):
        try:
            f = open(self.graph_file,"r",encoding='ISO-8859-16')
        except Exception:
            print("File does not exist")
            shutil.rmtree(self.get_new_folder_path(), ignore_errors=True)
            sys.exit(1)
        new_data_file = []
        data_file = f.readlines()
        field_row = data_file[0].replace('"', "")
        field_row = field_row.strip().replace(",", ";").split(";")
        counter, start, end = 0, 0, 0
        for field in field_row:
            if (field == "VarName"):
                start = counter
            elif (field == "Time_ms"):
                end = counter
            counter = counter + 1
        for row in data_file:
            record = row.replace('"', "").strip()
            record = record.replace(",", ";").split(";")[start:end+1]
            new_data_file.append(record)
        return new_data_file

    def chart_csv_update(self):
        plc_date = self.get_general_info()[0]
        new_workbook = Workbook()
        sheet = new_workbook.active
        data_file = self.get_graph_data()
        counter =0
        for row in data_file:
            time_difference = self.diff_times(plc_date,row[1])
            if(counter == 0):
                sheet.append(row)
                counter = counter +1
            elif (time_difference and row[0].find("$") == -1 ):
                row[2] = float(row[2])
                sheet.append(row)
        new_path = self.get_new_folder_path() + "/" + self.graph_file_name + "-" + self.get_process_num() + ".xlsx"
        new_workbook.save(new_path)
        style = StyleExcel(new_path)
        new_path = new_path.replace("/","\\")
        print(new_path)
        #style.run()
        injection = MacroInjection.MacroInjection(new_path)
        injection.run()


    def word_to_pdf(self):
        word_file = self.get_new_folder_path() + "/" + self.plc_file_name + "-" + self.get_process_num() + ".docx"
        pdf_file = self.get_new_folder_path() + "/" + self.plc_file_name + "-" + self.get_process_num() + ".pdf"
        word_file = word_file.replace("/","\\")
        pdf_file = pdf_file.replace("/","\\")
        word = win32com.client.Dispatch('Word.Application')
        print(word_file)
        print(pdf_file)
        doc = word.Documents.Open(word_file)
        doc.SaveAs(pdf_file, FileFormat = 17)
        doc.Close()
        word.Quit()


    def delete_unneeded(self):
        ext_files = [".xlsx",".pdf"]
        for file in os.listdir(self.get_new_folder_path()):
            if (os.path.splitext(file)[1] not in ext_files):
                os.remove(self.get_new_folder_path() + "/" + file)

    def killProcess(self):
        try:
            os.system("taskkill /f /im  EXCEL.EXE")
        except Exception:
            pass

    def run(self):
        try:
            self.killProcess()
            self.create_new_folder()
            self.modify_datafilewrite()
            self.chart_csv_update()
            self.write_to_word()
            self.word_to_pdf()
            self.delete_unneeded()
        except Exception:
            self.killProcess()



if __name__ == "__main__":
    x = Modifcation("")
    #x.modify_datafilewrite()
    #print("yes")
    x.write_to_word()